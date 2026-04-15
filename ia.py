"""
ia.py — Classificação de processos com múltiplos provedores de IA
Suporta: Anthropic (Claude), OpenAI (GPT/o-series), Google (Gemini),
         Groq (Llama/Mixtral), Mistral AI, DeepSeek, xAI (Grok)
"""
import re, json, time, os

# ══════════════════════════════════════════════════════════════
# CATÁLOGO DE MODELOS
# ══════════════════════════════════════════════════════════════
MODELO_PADRAO = "claude-haiku-4-5-20251001"

# Modelo econômico padrão por provedor (usado quando nenhum modelo é selecionado)
MODELO_PADRAO_POR_PROVIDER = {
    "anthropic": "claude-haiku-4-5-20251001",
    "openai":    "gpt-4o",
    "google":    "gemini-2.0-flash",
    "groq":      "llama-3.3-70b-versatile",
    "mistral":   "mistral-small-latest",
    "deepseek":  "deepseek-chat",
    "xai":       "grok-2",
}

# (label_grupo, provider_id, [(model_id, label_exibição), ...])
MODELOS_POR_PROVEDOR = [
    ("🟠 Anthropic — Claude", "anthropic", [
        ("claude-opus-4-6",            "Claude Opus 4.6 (Mais Poderoso)"),
        ("claude-sonnet-4-6",          "Claude Sonnet 4.6 (Recomendado)"),
        ("claude-haiku-4-5-20251001",  "Claude Haiku 4.5 (Rápido / Econômico)"),
        ("claude-3-5-sonnet-20241022", "Claude 3.5 Sonnet"),
        ("claude-3-5-haiku-20241022",  "Claude 3.5 Haiku"),
    ]),
    ("🟢 OpenAI — GPT / o-series", "openai", [
        ("gpt-4o",       "GPT-4o"),
        ("gpt-4o-mini",  "GPT-4o mini"),
        ("gpt-4-turbo",  "GPT-4 Turbo"),
        ("o1",           "o1 (Raciocínio)"),
        ("o1-mini",      "o1-mini (Raciocínio)"),
        ("o3-mini",      "o3-mini (Raciocínio)"),
    ]),
    ("🔵 Google — Gemini", "google", [
        ("gemini-2.0-flash",              "Gemini 2.0 Flash"),
        ("gemini-2.0-flash-thinking-exp", "Gemini 2.0 Flash Thinking"),
        ("gemini-1.5-pro",                "Gemini 1.5 Pro"),
        ("gemini-1.5-flash",              "Gemini 1.5 Flash"),
    ]),
    ("⚡ Meta — Llama / Mixtral (via Groq)", "groq", [
        ("llama-3.3-70b-versatile", "Llama 3.3 70B"),
        ("llama-3.1-8b-instant",    "Llama 3.1 8B"),
        ("mixtral-8x7b-32768",      "Mixtral 8x7B"),
    ]),
    ("🔷 Mistral AI", "mistral", [
        ("mistral-large-latest", "Mistral Large"),
        ("mistral-small-latest", "Mistral Small"),
        ("codestral-latest",     "Codestral"),
    ]),
    ("🐋 DeepSeek", "deepseek", [
        ("deepseek-chat",      "DeepSeek V3"),
        ("deepseek-reasoner",  "DeepSeek R1 (Reasoner)"),
    ]),
    ("🌟 xAI — Grok", "xai", [
        ("grok-2",  "Grok-2"),
        ("grok-3",  "Grok-3"),
    ]),
]

# Mapa rápido model_id → provider_id
_MODELO_PARA_PROVIDER = {
    mid: prov
    for _, prov, modelos in MODELOS_POR_PROVEDOR
    for mid, _ in modelos
}

# ══════════════════════════════════════════════════════════════
# FALLBACK DE MODELOS (Haiku/mini → Sonnet/4o quando SEM PARECER CONCLUSIVO)
# ══════════════════════════════════════════════════════════════
_FALLBACK_MODELOS = {
    "claude-haiku-4-5-20251001":  "claude-sonnet-4-6",
    "claude-3-5-haiku-20241022":  "claude-3-5-sonnet-20241022",
    "gpt-4o-mini":                "gpt-4o",
}

def _modelo_fallback(model):
    return _FALLBACK_MODELOS.get(model)


# ══════════════════════════════════════════════════════════════
# DETECÇÃO DE CRÉDITO ESGOTADO
# ══════════════════════════════════════════════════════════════
class CreditoEsgotadoError(Exception):
    """Crédito / quota esgotado no provedor de IA."""

_KEYWORDS_CREDITO = (
    "insufficient_balance", "insufficient balance", "credit balance",
    "quota_exceeded", "quota exceeded", "billing", "payment required",
    "your credit", "no credits", "out of credits", "credit limit",
    "account balance", "exceeded your", "you've exceeded",
    "resource_exhausted",
)

def _e_erro_credito(e):
    msg = str(e).lower()
    return any(k in msg for k in _KEYWORDS_CREDITO)


# ══════════════════════════════════════════════════════════════
# CACHE DE RESULTADOS (output/cache_ia.json)
# ══════════════════════════════════════════════════════════════
_CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output", "cache_ia.json")
_cache_ia: dict = {}
_cache_carregado = False

def _carregar_cache():
    global _cache_ia, _cache_carregado
    if _cache_carregado:
        return
    try:
        if os.path.exists(_CACHE_PATH):
            with open(_CACHE_PATH, encoding="utf-8") as f:
                _cache_ia = json.load(f)
    except Exception:
        _cache_ia = {}
    _cache_carregado = True

def _salvar_cache():
    try:
        os.makedirs(os.path.dirname(_CACHE_PATH), exist_ok=True)
        with open(_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump(_cache_ia, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def limpar_cache():
    global _cache_ia, _cache_carregado
    _cache_ia = {}
    _cache_carregado = True
    _salvar_cache()

def _chave_cache(numero):
    return re.sub(r"[^0-9]", "", numero)


# ══════════════════════════════════════════════════════════════
# PROMPTS ATUALIZADOS
# ══════════════════════════════════════════════════════════════
PROMPT_SISTEMA = """Você é um classificador jurídico especializado em processos das Turmas Recursais do TJAM.
Você analisa processos em que um advogado defende consumidores contra bancos/empresas.

═══ HIERARQUIA DOS DOCUMENTOS ═══
  1. ACÓRDÃO DE MÉRITO / SENTENÇA DE MÉRITO — documento principal (base para DECISAO e VALOR).
  2. ACÓRDÃO DE EMBARGOS / SENTENÇA DE EMBARGOS — pode MODIFICAR o mérito. Se acolhidos, use o resultado dos embargos.
  3. SENTENÇA DO 1º GRAU — contexto adicional.
  4. PETIÇÃO INICIAL — indispensável para MATÉRIA e valores de repetição de indébito.

  ⚠️ Se ACÓRDÃO DE MÉRITO e SENTENÇA DE MÉRITO estiverem "(não extraído)", DECISAO = "SEM PARECER CONCLUSIVO".

═══ REGRAS DE DECISÃO ═══
Valores válidos para DECISAO:
  • FAVORÁVEL | DESFAVORÁVEL | SENTENÇA ANULADA | EXTINTO SEM MÉRITO | SEM PARECER CONCLUSIVO

PASSO 0 — Identifique o advogado do AUTOR (subscrição da petição ou cabeçalho do acórdão).
PASSO 1 — Identifique o RECORRENTE. O consumidor pode ser RECORRENTE ou RECORRIDO — leia quem é quem.
PASSO 2 — Resultado do julgamento:
  "DAR PROVIMENTO" → RECORRENTE ganhou | "NEGAR PROVIMENTO" → RECORRENTE perdeu
  "ANULAR/CASSAR SENTENÇA" → SENTENÇA ANULADA | "EXTINTO" sem mérito → EXTINTO SEM MÉRITO
PASSO 3 — Combine (DECISAO reflete SEMPRE a perspectiva do consumidor):
  consumidor é RECORRENTE + DAR PROVIMENTO    → FAVORÁVEL
  consumidor é RECORRENTE + NEGAR PROVIMENTO  → DESFAVORÁVEL
  banco/réu  é RECORRENTE + DAR PROVIMENTO    → DESFAVORÁVEL
  banco/réu  é RECORRENTE + NEGAR PROVIMENTO  → FAVORÁVEL

  ⚠️ ATENÇÃO: Mesmo que o banco seja o recorrente e ganhe o recurso, se o resultado
  final prejudica o consumidor, DECISAO = DESFAVORÁVEL. O STATUS é sempre do
  ponto de vista do consumidor, nunca do banco.

⚠️ VERIFICAÇÃO: DECISAO deve ser CONSISTENTE com o RACIOCINIO. Revise se houver contradição.

═══ REGRAS DE MATÉRIA ═══
Fontes (em ordem): [ARQUIVO: ...] no texto → ementa/cabeçalho → fatos da petição inicial.
PRIORIDADE: petição inicial é a fonte principal para identificar a matéria.

Use EXATAMENTE um dos códigos abaixo. Para cada um: leia o IDENTIFICAR e o NÃO CONFUNDIR antes de decidir.

  AD_DEPOSITANTE      — Desconto de Adiantamento a Depositante. IDENTIFICAR quando: a petição contesta desconto em conta corrente sob a rubrica "Adiantamento a Depositante" ou tarifa equivalente por falta de contratação. NÃO CONFUNDIR com TARIFA_IND: AD_DEPOSITANTE refere-se estritamente a esta rubrica específica.

  ANP                 — Negativação sem notificação prévia. IDENTIFICAR quando: a ação questiona inclusão do nome no SPC/Serasa sem notificação prévia por correspondência postal. NÃO CONFUNDIR com ANP_NN (que é plataforma de negociação, não negativação) nem ANP_PROTESTO (que envolve cartório).

  ANP_NN              — Cobrança em plataforma de negociação sem notificação. IDENTIFICAR quando: questiona inclusão de dívida em plataformas como Serasa Limpa Nome sem aviso prévio, sem gerar restrição de crédito clássica. NÃO CONFUNDIR com ANP (negativação real no SPC/Serasa) nem SLN (que foca na prescrição da dívida).

  ANP_PROTESTO        — Protesto sem notificação prévia. IDENTIFICAR quando: questiona inclusão de restrição via Protesto em Cartório sem aviso prévio. NÃO CONFUNDIR com ANP: aqui envolve cartório de protesto, não SPC/Serasa.

  ANUIDADE            — Anuidade de cartão de crédito não autorizada. IDENTIFICAR quando: a petição contesta desconto de anuidade diretamente na conta ou fatura do cartão. NÃO CONFUNDIR com GASTO_C_CRED: ANUIDADE é a tarifa de manutenção do cartão, não débitos de compras.

  BLOQUEIO_APP_MOTORISTA — Bloqueio indevido de app de motorista. IDENTIFICAR quando: o autor é motorista (Uber, 99, etc.) e processa a plataforma por bloqueio/banimento unilateral sem aviso ou provas. NÃO CONFUNDIR com ações trabalhistas: o foco é reativação da conta e danos pelo bloqueio.

  BX_ANT_FIN          — Baixa antecipada de financiamento. IDENTIFICAR quando: a ação questiona lançamentos na conta com a sigla "BX.ANT.FIN/EMP" ou liquidação antecipada não solicitada de empréstimo. NÃO CONFUNDIR com PARC_CRED_PESS: BX_ANT_FIN é uma liquidação/baixa antecipada, não a parcela normal do empréstimo.

  CESTA               — Pacote de serviços bancários cobrado indevidamente. IDENTIFICAR quando: contesta descontos de pacotes mensais (ex: Cesta B. Expresso, Cesta Fácil) por ausência de contrato assinado. NÃO CONFUNDIR com TARIFA_IND: CESTA é mensalidade do pacote; TARIFA_IND são tarifas avulsas por uso.

  COBRANCA_IND        — Cobrança indevida genérica. IDENTIFICAR quando: a ação contesta um desconto sem rubrica específica mapeada nas demais matérias deste sistema. NÃO CONFUNDIR com OUTRO: COBRANCA_IND ainda é especificamente sobre cobrança/desconto financeiro indevido, apenas sem rubrica catalogada.

  COI_BOLETO          — Cobrança indevida via boleto por recuperadoras/empresas de cobrança. IDENTIFICAR quando: a ação contesta recebimento de boleto bancário enviado por empresa recuperadora de crédito (Recovery, Ipanema, Boa Vista, PEFISA, etc.) ou por qualquer terceiro cobrando dívida prescrita ou indevida via boleto — mesmo que o nome da recuperadora não apareça explicitamente, se o fato central é "recebi boleto cobrando dívida que não devo / prescrita", use COI_BOLETO. NÃO CONFUNDIR com SLN: COI_BOLETO é envio de boleto físico ou digital; SLN é negativação na plataforma Serasa Limpa Nome. NÃO CONFUNDIR com COBRANCA_IND: COBRANCA_IND é desconto direto em conta corrente/benefício sem boleto externo; COI_BOLETO sempre envolve boleto enviado ao consumidor por terceiro cobrador.

  COI_SIND            — Desconto sindical não autorizado no benefício INSS. IDENTIFICAR quando: contesta descontos em benefício previdenciário ou contracheque EXCLUSIVAMENTE para sindicatos ou associações (ex: ASSEJUS, FENASPS) sem autorização do trabalhador. NÃO CONFUNDIR com RMC, PARC_CRED_PESS ou REFINANCIAMENTO_IND: COI_SIND é exclusivo de sindicato/associação — se o desconto é de parcela de empréstimo consignado no INSS (mesmo que não autorizado), use PARC_CRED_PESS ou REFINANCIAMENTO_IND conforme o caso.

  EMISSAO_EXTRATO     — ⚠️ SIGLA DESCONTINUADA NESTE SISTEMA. Use TARIFA_IND para cobranças de emissão de extrato bancário.

  ENERGIA_TARIFAS     — Tarifas indevidas em fatura de energia. IDENTIFICAR quando: contesta rubricas como "Encargos", "Perdas" e "Outros" na fatura de energia elétrica da concessionária. NÃO CONFUNDIR com TOI (que é multa por inspeção de suposto desvio) nem ESPECIFICA_ENERGIA (narrativa factual única, danos singulares).

  ESPECIFICA          — Ação cível/consumidor customizada com narrativa única. IDENTIFICAR quando: a petição tem narrativa extensa e personalizada descrevendo fatos únicos do cliente (situação concreta, histórico detalhado, prejudicados específicos) que NÃO seguem o modelo padronizado do escritório. Típico: petição de 8+ páginas com situação factual singular. NÃO CONFUNDIR com COBRANCA_IND: COBRANCA_IND é cobrança financeira genérica; ESPECIFICA tem narrativa rica e individual mesmo que envolva um produto conhecido. NÃO CONFUNDIR com OUTRO: ESPECIFICA ainda pertence ao universo de direito do consumidor/bancário, só não cabe no modelo de massa.

  ESPECIFICA_ENERGIA  — Ação complexa de energia com narrativa única. IDENTIFICAR quando: é contra concessionária de energia com narrativa factual singular (ex: queima de aparelhos), fora dos modelos. NÃO CONFUNDIR com ENERGIA_TARIFAS (tarifas padronizadas de fatura) nem TOI (multa de inspeção).

  EXIBICAO            — Ação de exibição de documentos. IDENTIFICAR quando: é ação cautelar ou de provas pedindo a exibição de um contrato bancário. NÃO CONFUNDIR com ações indenizatórias: o fim principal de EXIBICAO é obter o documento, não indenização.

  EXTRATO_MOVIMENTO   — ⚠️ SIGLA DESCONTINUADA NESTE SISTEMA. Use TARIFA_IND para cobranças de extrato de movimentação bancária.

  GASTO_C_CRED        — Gastos de cartão debitados indevidamente na conta corrente. IDENTIFICAR quando: questiona débitos automáticos indevidos na conta sob a rubrica "Gastos Cartão de Crédito". NÃO CONFUNDIR com PARC_AUTOMATICO: GASTO_C_CRED é o débito do gasto direto; PARC_AUTOMATICO é o parcelamento imposto do saldo devedor da fatura.

  GOLPE_PIX           — Fraude / golpe via Pix. IDENTIFICAR quando: pede indenização por falha de segurança do banco ao não bloquear transferências atípicas via Pix ou por engenharia social no aplicativo. NÃO CONFUNDIR com COI_BOLETO: GOLPE_PIX é invasão/fraude no app bancário; COI_BOLETO é envio de boleto falso por recuperadora.

  INCORP_BANCOS       — Rubrica "Incorporação Bancos". IDENTIFICAR quando: contesta cobrança no extrato exatamente sob a rubrica "INCORPORAÇÃO BANCOS" (comum no Bradesco). NÃO CONFUNDIR com TARIFA_IND: INCORP_BANCOS tem nomenclatura própria e tese específica definida pelo escritório.

  INV_FACIL           — Investimento compulsório não solicitado. IDENTIFICAR quando: contesta direcionamento automático de saldo da conta para aplicações (ex: Invest Fácil, CDB automático) bloqueando o dinheiro sem autorização. NÃO CONFUNDIR com CESTA: INV_FACIL não é taxa de serviço, mas retenção do dinheiro em fundo de investimento.

  JUROS_ABUSIVOS      — Juros acima da taxa média do BACEN. IDENTIFICAR quando: pede redução da taxa de juros do contrato alegando ser superior à taxa média de mercado do BACEN. NÃO CONFUNDIR com JUROS_NC: JUROS_ABUSIVOS ataca a taxa que está no contrato; JUROS_NC ataca cobrança acima do que foi contratado.

  JUROS_ABUSIVOS+COI  — Juros abusivos + Comissão de Originação. IDENTIFICAR quando: além dos juros abusivos, a petição também contesta a cobrança de COI (Comissão de Originação / Comissão de Intermediação).

  JUROS_ABUSIVOS+COI+CN — Juros abusivos + COI + Capitalização. IDENTIFICAR quando: a petição contesta simultaneamente juros abusivos, COI e capitalização de juros (juros sobre juros) no mesmo contrato.

  JUROS_NC            — Juros não contratados. IDENTIFICAR quando: alega que o banco cobra na prática uma taxa maior do que a escrita no contrato assinado. NÃO CONFUNDIR com JUROS_ABUSIVOS: JUROS_NC é descumprimento do contrato; JUROS_ABUSIVOS é abusividade da cláusula pactuada.

  MORA                — Encargos de mora não reconhecidos. IDENTIFICAR quando: contesta descontos no extrato sob rubricas de penalidade por atraso (MORA CRED PESS, ENCARGOS DE MORA). NÃO CONFUNDIR com PARC_CRED_PESS: MORA é a multa/juro do atraso, não a parcela principal do empréstimo.

  MORA_CEL            — Mora em contrato de celular/operadora. IDENTIFICAR quando: a mora ou encargo indevido refere-se a contrato com operadora de telefonia celular.

  OUTRO               — Matéria completamente fora do sistema. IDENTIFICAR quando: a ação NÃO envolve qualquer produto bancário, financeiro, de telefonia ou energia — é matéria completamente estranha ao portfólio do escritório. Use com extrema parcimônia. NÃO CONFUNDIR com COBRANCA_IND (qualquer cobrança financeira indevida sem rubrica específica — prefira COBRANCA_IND a OUTRO) nem com ESPECIFICA (ação de consumidor/bancário com narrativa única — prefira ESPECIFICA a OUTRO). Se o produto ou serviço contestado puder ser identificado mesmo que de forma vaga, use a matéria correspondente, nunca OUTRO.

  PARC_AUTOMATICO     — Parcelamento automático de fatura do cartão. IDENTIFICAR quando: contesta parcelamento do saldo da fatura do cartão de crédito imposto pelo banco após pagamento mínimo. NÃO CONFUNDIR com RMC: PARC_AUTOMATICO ocorre em cartão de crédito comum; RMC é cartão consignado com desconto em folha.

  PARC_CRED_PESS      — Empréstimo pessoal não contratado. IDENTIFICAR quando: contesta descontos na conta sob a rubrica "Parcela Crédito Pessoal" ou "PARC CRED PESS" sem contrato assinado. NÃO CONFUNDIR com REFINANCIAMENTO_IND: PARC_CRED_PESS é empréstimo novo não reconhecido; REFINANCIAMENTO_IND é renegociação não autorizada de empréstimo que o autor reconhece.

  RCC                 — Reserva de Cartão Consignado. IDENTIFICAR quando: a petição menciona explicitamente a sigla "RCC" ou "Reserva de Cartão Consignado". Tese: produto vendido como empréstimo mas era cartão consignado gerando dívida rotativa. NÃO CONFUNDIR com RMC: a diferença é a rubrica/nomenclatura usada na petição.

  REFINANCIAMENTO_IND — Refinanciamento não autorizado de empréstimo. IDENTIFICAR quando: o autor reconhece o empréstimo original, mas impugna o refinanciamento/averbação que gerou novo prazo ou desconto sem autorização. NÃO CONFUNDIR com PARC_CRED_PESS: aqui o autor reconhece o contrato base, atacando apenas a renegociação unilateral.

  REFINANCIAMENTO_IND_ESPECIFICA — Refinanciamento não autorizado com narrativa complexa. IDENTIFICAR quando: é ação de refinanciamento não autorizado, mas com petição extensa e fatos únicos fora do modelo de massa. NÃO CONFUNDIR com REFINANCIAMENTO_IND: esta sigla é para casos altamente customizados do caso concreto.

  RMC                 — Reserva de Margem Consignável. IDENTIFICAR quando: a petição menciona explicitamente a sigla "RMC" ou "Reserva de Margem Consignável". Tese idêntica ao RCC (cartão vendido como empréstimo), mas a distinção é a nomenclatura usada na petição. NÃO CONFUNDIR com RCC.

  SAQUE_TERMINAL      — ⚠️ SIGLA DESCONTINUADA NESTE SISTEMA. Use TARIFA_IND para cobranças de saque em terminal (caixa eletrônico).

  SEGURO              — Seguro embutido em empréstimo/financiamento. IDENTIFICAR quando: contesta seguros (Prestamista, Proteção Financeira, Seguro de Vida, Seguro Prestamista) embutidos em contratos de empréstimo, crédito pessoal ou financiamento sem autorização do cliente. NÃO CONFUNDIR com SEGURO_CARTAO: se o seguro vem junto a contrato de empréstimo ou crédito pessoal use SEGURO; somente use SEGURO_CARTAO se o seguro aparece exclusivamente na fatura do cartão de crédito (Perda e Roubo, Fatura Protegida, Proteção Cartão).

  SEGURO_CARTAO       — Seguro cobrado na fatura do cartão de crédito. IDENTIFICAR quando: contesta seguros (Perda e Roubo, Fatura Protegida, Proteção Cartão) cobrados na fatura do cartão de crédito. NÃO CONFUNDIR com SEGURO: esta matéria é exclusiva de seguros atrelados ao cartão, não a empréstimos.

  SLN                 — Dívida prescrita cobrada no Serasa Limpa Nome. IDENTIFICAR quando: pede inexigibilidade de dívida prescrita (mais de 5 anos) cobrada na plataforma que afeta o score de crédito. NÃO CONFUNDIR com ANP_NN: SLN foca na prescrição da dívida; ANP_NN foca na falta de aviso prévio.

  SVA                 — Serviço de Valor Adicionado de telefonia. IDENTIFICAR quando: questiona inclusão de SVA não solicitado em faturas de telefonia/consumo (venda casada em operadora). NÃO CONFUNDIR com CESTA: SVA é cobrado por operadoras de telefonia; CESTA é pacote de serviços bancários.

  TARIFA_CAD          — Tarifa de Cadastro em financiamento. IDENTIFICAR quando: questiona a legalidade da Tarifa de Cadastro (TC) no início de contratos de financiamento (ex: veículos). NÃO CONFUNDIR com outras tarifas: TARIFA_CAD é exclusiva de contratos de financiamento.

  TARIFA_IND          — Tarifa bancária avulsa indevida (categoria principal). IDENTIFICAR quando: a ação contesta qualquer tarifa bancária cobrada individualmente sem autorização contratual. INCLUI: saque em terminal/caixa eletrônico, emissão de extrato, extrato de movimentação, TED/DOC, tarifa de manutenção de conta, ou qualquer outro serviço bancário avulso. É a matéria PADRÃO para tarifas bancárias — se a rubrica não tiver sigla própria neste sistema, use TARIFA_IND. NÃO CONFUNDIR com CESTA (mensalidade de pacote completo de serviços) nem com AD_DEPOSITANTE (rubrica específica de adiantamento) nem com INCORP_BANCOS (rubrica específica do Bradesco) nem com TARIFA_CAD (tarifa de cadastro de financiamento). ATENÇÃO: mesmo que o documento mencione "emissão de extrato", "saque no terminal" ou "extrato de movimentação" como rubrica específica, USE TARIFA_IND — essas sub-rubricas não têm sigla separada neste sistema.

  TIT_CAP             — Título de Capitalização não contratado. IDENTIFICAR quando: contesta desconto automático de Títulos de Capitalização não solicitados na conta corrente. NÃO CONFUNDIR com INV_FACIL: TIT_CAP é produto de capitalização com sorteios; INV_FACIL retém saldo em aplicações/investimentos de liquidez.

  TOI                 — Termo de Ocorrência e Inspeção. IDENTIFICAR quando: ataca multa imposta por concessionária de água ou energia após vistoria de suposto desvio/fraude. NÃO CONFUNDIR com ENERGIA_TARIFAS: TOI é penalidade de inspeção; ENERGIA_TARIFAS são rubricas de cobrança comuns da fatura mensal.

  VIDA_PREV           — Desconto de seguro de vida ou previdência. IDENTIFICAR quando: contesta descontos no extrato bancário a título de "Bradesco Vida e Previdência" ou produto similar debitado sem contratação. NÃO CONFUNDIR com SEGURO: VIDA_PREV foca especificamente em produtos de previdência/vida debitados avulsamente em conta.

⚠️ REGRAS ESPECIAIS DE MATÉRIA:
  • Se DECISAO = "SEM PARECER CONCLUSIVO" → classifique MATERIA com base na PETIÇÃO INICIAL
    se ela estiver disponível e permitir identificar o produto/serviço contestado.
    Use "OUTRO" apenas se a petição inicial também estiver ausente ou não permitir identificar a matéria.
  • Se DECISAO = "EXTINTO SEM MÉRITO" → classifique MATERIA com base na petição inicial
    ou acórdão disponível, igual às demais decisões. Use "OUTRO" apenas se nenhum documento
    permitir identificar o produto ou serviço contestado.
  • Se DECISAO = "SENTENÇA ANULADA" → use a matéria da petição inicial se disponível.

═══ REGRAS DE VALORES DE CONDENAÇÃO ═══
Procure em: acórdão/embargos de mérito → sentença → petição inicial.
Use o valor do acórdão se ele reformou a sentença.
Preencha DANO_MATERIAL e DANO_MORAL separadamente. Formato obrigatório: R$ #.##0,00 (ex: R$ 3.000,00).
Deixe ambos vazios se DESFAVORÁVEL, EXTINTO ou SEM PARECER CONCLUSIVO.

⚠️ Se DECISAO = FAVORÁVEL mas não há condenação monetária:
  — Descreva a tutela obtida no RACIOCINIO (inexigibilidade, obrigação de fazer, suspensão de desconto etc.)
  — Deixe DANO_MORAL e DANO_MATERIAL vazios, mas NUNCA deixe o RACIOCINIO vago sobre o que o consumidor ganhou.
⚠️ Se DECISAO = EXTINTO SEM MÉRITO ou SEM PARECER CONCLUSIVO:
  — DANO_MORAL e DANO_MATERIAL devem estar vazios obrigatoriamente.

ATENÇÃO: Responda SOMENTE com o bloco JSON, sem texto antes ou depois, sem markdown.
"""

PROMPT_USER = """Classifique o processo abaixo e retorne APENAS o JSON.

NÚMERO: {numero}
TIPO DE DECISÃO: {tipo}
TURMA/VARA: {turma_vara}
RELATOR/JUIZ: {relator_juiz}
AUTOR/RECORRENTE: {recorrente}
RÉU/RECORRIDO: {recorrido}
ASSUNTO: {assunto}
ADVOGADO DO AUTOR: {nome_advogado_hint}

══ {label_merito} (BASE PRINCIPAL — use para DECISÃO e VALOR) ══
{texto_principal}

══ {label_embargos_principal} (se presente: verifique se é Embargos de Declaração e se acolheu alterações) ══
{texto_embargos_principal}
{aviso_dois_documentos}

══ SENTENÇA DO 1º GRAU (contexto) ══
{texto_sentenca}

══ EMBARGOS DE DECLARAÇÃO DA SENTENÇA (verifique alterações) ══
{texto_sentenca_embargos}

══ PETIÇÃO INICIAL (fonte para MATÉRIA e valores de repetição de indébito) ══
{texto_peticao}

Retorne SOMENTE este JSON (sem markdown):
{{
  "ADVOGADO": "Nome do advogado do autor identificado nos documentos",
  "DECISAO": "FAVORÁVEL, DESFAVORÁVEL, SENTENÇA ANULADA, EXTINTO SEM MÉRITO ou SEM PARECER CONCLUSIVO",
  "MATERIA": "SIGLA_DA_MATERIA",
  "DANO_MATERIAL": "Valor do dano material / repetição de indébito / devolução condenada, somente se FAVORÁVEL. Formato obrigatório: R$ #.##0,00 (ex: R$ 1.500,00). Vazio se não houver ou se DESFAVORÁVEL/EXTINTO/SEM PARECER.",
  "DANO_MORAL": "Valor do dano moral condenado, somente se FAVORÁVEL. Formato obrigatório: R$ #.##0,00 (ex: R$ 3.000,00). Vazio se não houver ou se DESFAVORÁVEL/EXTINTO/SEM PARECER.",
  "RACIOCINIO": "3 frases obrigatórias: (1) Quem é o recorrente e qual foi o resultado — se FAVORÁVEL, descreva QUAL benefício concreto o consumidor obteve (ex: danos morais de R$ X, inexigibilidade de débito de R$ Y, suspensão de desconto indevido); se DESFAVORÁVEL, descreva o que foi negado. (2) Houve embargos de declaração? Foram acolhidos? O que mudou na decisão ou nos valores? (3) Discriminação dos valores condenados: danos morais + materiais + repetição de indébito + outros. Se FAVORÁVEL sem condenação monetária, descreva a tutela obtida (inexigibilidade, obrigação de fazer etc.). NUNCA escreva apenas 'recurso foi provido' sem explicar o que isso significa para o consumidor."
}}"""


# ══════════════════════════════════════════════════════════════
# UTILITÁRIO — recorte de texto (início + fim para não perder o dispositivo)
# ══════════════════════════════════════════════════════════════
def _trim(texto, max_chars, tail=2500):
    """Mantém o início (partes/fatos) e o fim (dispositivo) do documento."""
    if not texto or len(texto) <= max_chars:
        return texto or ""
    head = max_chars - tail
    return texto[:head] + "\n[...trecho omitido...]\n" + texto[-tail:]


# ══════════════════════════════════════════════════════════════
# UTILITÁRIO — detectar provider
# ══════════════════════════════════════════════════════════════
def _detectar_provider(model):
    if model in _MODELO_PARA_PROVIDER:
        return _MODELO_PARA_PROVIDER[model]
    # fallback por prefixo
    if model.startswith("claude"):     return "anthropic"
    if model.startswith(("gpt-","o1","o3")): return "openai"
    if model.startswith("gemini"):     return "google"
    if model.startswith(("llama","mixtral")): return "groq"
    if model.startswith(("mistral","codestral")): return "mistral"
    if model.startswith("deepseek"):   return "deepseek"
    if model.startswith("grok"):       return "xai"
    return "anthropic"


# ══════════════════════════════════════════════════════════════
# FALLBACK DE PROVEDORES (crédito esgotado)
# ══════════════════════════════════════════════════════════════
def _carregar_api_keys():
    """Lê todas as chaves de API do config.ini."""
    import configparser
    cfg = configparser.ConfigParser()
    cfg.read(os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.ini"), encoding="utf-8")
    return {
        "anthropic": cfg.get("claude",    "api_key", fallback=""),
        "openai":    cfg.get("openai",    "api_key", fallback=""),
        "google":    cfg.get("google",    "api_key", fallback=""),
        "groq":      cfg.get("groq",      "api_key", fallback=""),
        "mistral":   cfg.get("mistral",   "api_key", fallback=""),
        "deepseek":  cfg.get("deepseek",  "api_key", fallback=""),
        "xai":       cfg.get("xai",       "api_key", fallback=""),
    }


def _sequencia_fallback(model_inicial, key_inicial):
    """
    Retorna lista de (model, api_key) a tentar em sequência quando crédito esgotar.
    Começa pelo modelo solicitado e percorre os demais provedores configurados.
    """
    todas = _carregar_api_keys()
    seq = [(model_inicial, key_inicial)]
    for _, prov, _ in MODELOS_POR_PROVEDOR:
        chave = todas.get(prov, "").strip()
        if not chave:
            continue
        modelo = MODELO_PADRAO_POR_PROVIDER[prov]
        if modelo != model_inicial:
            seq.append((modelo, chave))
    return seq


# ══════════════════════════════════════════════════════════════
# CHAMADA UNIFICADA À LLM
# ══════════════════════════════════════════════════════════════
def _chamar_llm(system_prompt, user_prompt, model, api_key, max_tokens=512):
    """Chama o provider correto e retorna o texto da resposta."""
    provider = _detectar_provider(model)

    # ── Anthropic ──────────────────────────────────────────────
    if provider == "anthropic":
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        is_v4 = any(model.startswith(p) for p in ("claude-opus-4", "claude-sonnet-4", "claude-haiku-4"))
        system_block = [{"type": "text", "text": system_prompt,
                          "cache_control": {"type": "ephemeral"}}]
        kwargs = dict(
            model=model,
            max_tokens=max_tokens,
            temperature=1 if is_v4 else 0,
            system=system_block,
            messages=[{"role": "user", "content": user_prompt}],
            extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
        )
        if is_v4:
            kwargs["thinking"] = {"type": "disabled"}
        texto = ""
        with client.messages.stream(**kwargs) as stream:
            for t in stream.text_stream:
                texto += t
        return texto

    # ── OpenAI-compatíveis (OpenAI, Groq, DeepSeek, xAI) ──────
    if provider in ("openai", "groq", "deepseek", "xai"):
        import openai
        base_urls = {
            "groq":     "https://api.groq.com/openai/v1",
            "deepseek": "https://api.deepseek.com",
            "xai":      "https://api.x.ai/v1",
        }
        client = openai.OpenAI(
            api_key=api_key,
            base_url=base_urls.get(provider),
        )
        is_reasoning = model in ("o1", "o1-mini", "o3-mini", "o1-pro", "deepseek-reasoner")
        if is_reasoning:
            msgs = [{"role": "user", "content": f"{system_prompt}\n\n{user_prompt}"}]
            kwargs = {"model": model, "max_completion_tokens": max_tokens, "messages": msgs}
        else:
            msgs = [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ]
            kwargs = {"model": model, "max_tokens": max_tokens, "temperature": 0, "messages": msgs}
            if provider == "openai":
                kwargs["response_format"] = {"type": "json_object"}
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message.content or ""

    # ── Google Gemini ──────────────────────────────────────────
    if provider == "google":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        modelo_obj = genai.GenerativeModel(
            model_name=model,
            system_instruction=system_prompt,
            generation_config={"temperature": 0, "max_output_tokens": max_tokens},
        )
        resp = modelo_obj.generate_content(user_prompt)
        return resp.text or ""

    # ── Mistral AI ────────────────────────────────────────────
    if provider == "mistral":
        from mistralai import Mistral
        client = Mistral(api_key=api_key)
        resp = client.chat.complete(
            model=model,
            temperature=0,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
        )
        return resp.choices[0].message.content or ""

    raise ValueError(f"Provider desconhecido para o modelo '{model}'")


# ══════════════════════════════════════════════════════════════
# EXTRAÇÃO DE PARTES
# ══════════════════════════════════════════════════════════════
def extrair_partes(texto):
    """Extrai recorrente, recorrido, autor, réu e assunto do texto da peça."""
    info = {"recorrente": "", "recorrido": "", "autor": "", "reu": "", "assunto": ""}
    if not texto:
        return info
        
    for linha in texto.splitlines():
        s  = linha.strip()
        sl = s.lower()
        if not info["recorrente"] and "recorrente" in sl and ":" in s and "recorrido" not in sl:
            v = s.split(":", 1)[-1].strip().rstrip(")")
            if 3 < len(v) < 150:
                info["recorrente"] = v
        if not info["recorrido"] and "recorrido" in sl and ":" in s:
            v = s.split(":", 1)[-1].strip().rstrip(")")
            if 3 < len(v) < 150:
                info["recorrido"] = v
        if not info["autor"] and "autor" in sl and ":" in s and "autori" not in sl:
            v = s.split(":", 1)[-1].strip()
            if 3 < len(v) < 150:
                info["autor"] = v
        if not info["reu"] and ("réu" in sl or "reu" in sl) and ":" in s:
            v = s.split(":", 1)[-1].strip()
            if 3 < len(v) < 150:
                info["reu"] = v
        if not info["assunto"] and any(sl.startswith(x) for x in ["assunto", "objeto"]) and ":" in s:
            v = s.split(":", 1)[-1].strip()
            if 3 < len(v) < 300:
                info["assunto"] = v
    return info


# ══════════════════════════════════════════════════════════════
# CLASSIFICADOR PRINCIPAL COM RETRY
# ══════════════════════════════════════════════════════════════
def classificar(numero, tipo, turma_vara, relator_juiz,
                partes, texto_principal, texto_sentenca, texto_peticao,
                api_key, log, model=None, nome_advogado="",
                texto_embargos_principal="", texto_sentenca_embargos=""):
    """Classifica o processo com a LLM escolhida. Possui mecanismo de retentativa."""
    model_usar = model if model else MODELO_PADRAO
    provider   = _detectar_provider(model_usar)

    if tipo == "ACÓRDÃO":
        label_merito            = "ACÓRDÃO DE MÉRITO (2º GRAU)"
        label_embargos_principal = "ACÓRDÃO DE EMBARGOS DE DECLARAÇÃO (2º GRAU)"
    else:
        label_merito            = "SENTENÇA DE MÉRITO (1º GRAU)"
        label_embargos_principal = "EMBARGOS DE DECLARAÇÃO DA SENTENÇA (1º GRAU)"

    nome_advogado_hint = nome_advogado.strip() if nome_advogado and nome_advogado.strip() else ""

    # Aviso quando há 2 documentos — pede à IA confirmar mérito vs embargos
    if texto_embargos_principal and texto_embargos_principal.strip() not in ("", "(não extraído)"):
        aviso = ("⚠️ Foram encontrados 2 documentos. O primeiro é presumido como MÉRITO e o segundo "
                 "como EMBARGOS DE DECLARAÇÃO. Confirme pelo conteúdo: se o segundo não for embargos, "
                 "trate ambos como acórdãos/sentenças de mérito e use o mais recente como base.")
    else:
        aviso = ""

    prompt = PROMPT_USER.format(
        numero                   = numero,
        tipo                     = tipo,
        turma_vara               = turma_vara,
        relator_juiz             = relator_juiz,
        recorrente               = partes.get("recorrente") or partes.get("autor", ""),
        recorrido                = partes.get("recorrido")  or partes.get("reu", ""),
        assunto                  = partes.get("assunto", ""),
        nome_advogado_hint       = nome_advogado_hint,
        label_merito             = label_merito,
        label_embargos_principal = label_embargos_principal,
        aviso_dois_documentos    = aviso,
        # Recorte inteligente: início (partes/fatos) + fim (dispositivo)
        texto_principal          = _trim(texto_principal,          30000, tail=6000) or "(não extraído)",
        texto_embargos_principal = _trim(texto_embargos_principal, 15000, tail=4000) or "(não extraído)",
        texto_sentenca           = _trim(texto_sentenca,           15000, tail=4000) or "(não extraída)",
        texto_sentenca_embargos  = _trim(texto_sentenca_embargos,  10000, tail=3000) or "(não extraída)",
        texto_peticao            = _trim(texto_peticao,            20000, tail=0)    or "(não extraída)",
    )

    # ── Cache ──────────────────────────────────────────────────
    _carregar_cache()
    chave = _chave_cache(numero)
    if chave in _cache_ia:
        log(f"   💾 Resultado do cache (análise anterior).")
        return _cache_ia[chave]

    log(f"   🤖 [{provider.upper()}] {model_usar}...")

    def _chamar_e_parsear(model, prompt_user, key):
        """Chama a LLM e retorna o dict parseado. Lança CreditoEsgotadoError se sem crédito."""
        for tentativa in range(3):
            try:
                texto_resp = _chamar_llm(PROMPT_SISTEMA, prompt_user, model, key, max_tokens=2000)
                limpo = re.sub(r'^```(?:json)?\s*|\s*```\s*$', '', texto_resp).strip()
                try:
                    return json.loads(limpo)
                except json.JSONDecodeError:
                    m = re.search(r'\{[\s\S]*\}', limpo)
                    if m:
                        return json.loads(m.group())
            except CreditoEsgotadoError:
                raise
            except Exception as e:
                if _e_erro_credito(e):
                    raise CreditoEsgotadoError(str(e)) from e
                if tentativa < 2:
                    log(f"   ⚠️ Falha na API ({e}). Tentando novamente em 3s...")
                    time.sleep(3)
                else:
                    log(f"   ❌ Erro final na API após 3 tentativas: {e}")
                    return {"RACIOCINIO": f"Erro na API da IA: {e}"}
        return {}

    # ── Sequência de fallback por crédito ────────────────────
    seq = _sequencia_fallback(model_usar, api_key)
    resultado = {}
    key_usar = api_key
    for idx_seq, (m_tent, k_tent) in enumerate(seq):
        if idx_seq > 0:
            prov = _detectar_provider(m_tent)
            log(f"   🔄 Sem crédito — tentando {m_tent} [{prov.upper()}]...")
        try:
            resultado = _chamar_e_parsear(m_tent, prompt, k_tent)
            model_usar = m_tent  # atualiza para o fallback do SEM PARECER
            key_usar   = k_tent  # salva a chave que funcionou
            break
        except CreditoEsgotadoError:
            log(f"   ⚠️ Crédito esgotado em {_detectar_provider(m_tent).upper()}.")
    if not resultado:
        log("   ❌ Todos os provedores sem crédito ou indisponíveis.")
        return {}

    # ── Fallback automático: Haiku/mini → Sonnet/4o ──────────
    if resultado and resultado.get("DECISAO") == "SEM PARECER CONCLUSIVO":
        fallback = _modelo_fallback(model_usar)
        if fallback:
            log(f"   ⚠️ SEM PARECER CONCLUSIVO — retentando com {fallback}...")
            try:
                resultado_fb = _chamar_e_parsear(fallback, prompt, key_usar)
            except CreditoEsgotadoError:
                log(f"   ⚠️ Crédito esgotado no fallback {fallback}.")
                resultado_fb = None
            if resultado_fb:
                resultado = resultado_fb

    # ── Salva no cache ────────────────────────────────────────
    if resultado and "RACIOCINIO" in resultado and not resultado.get("RACIOCINIO", "").startswith("Erro"):
        _cache_ia[chave] = resultado
        _salvar_cache()

    return resultado or {}


# ══════════════════════════════════════════════════════════════
# RELATÓRIO ANALÍTICO
# ══════════════════════════════════════════════════════════════
_PROMPT_RELATORIO_SISTEMA = """Você é um analista jurídico das Turmas Recursais do TJAM.
Produza relatórios objetivos, estratégicos e consultáveis — sem enrolação.
Regras obrigatórias: linguagem direta e técnica; toda afirmação sobre caso específico deve citar o número do processo; sem repetições entre seções; máximo de 3 linhas por ponto."""

_PROMPT_RELATORIO_USER = """Analise os {n} processos abaixo e gere um relatório estratégico conciso.

DADOS DOS PROCESSOS:
{tabela}

Estrutura OBRIGATÓRIA — use ## para cada título de seção:

## 1. PANORAMA GERAL
Linha única de estatísticas: total | favoráveis (%) | desfavoráveis (%) | sem parecer (%) | extintos/outros (%).
Em seguida, 1 parágrafo curto com a principal conclusão estratégica do lote.

## 2. ANÁLISE POR RELATOR / JUIZ
Para cada magistrado com 2 ou mais processos, uma linha no formato:
▸ [Nome] — N proc. | Favoráveis: X (Y%) | Padrão: [tendência em 1 frase] | Ref.: [nº dos processos]

## 3. ANÁLISE POR MATÉRIA
Para cada matéria com 2 ou mais processos, uma linha no formato:
▸ [Matéria] — N proc. | Favoráveis: X (Y%) | Valor médio condenação: R$ Z | Ref.: [nº dos processos]

## 4. DESTAQUES DE CONDENAÇÃO
Apenas os 3 maiores valores obtidos: número do processo | valor dano moral | valor dano material.
Valor médio geral de dano moral (favoráveis). Valor médio geral de dano material (favoráveis).

## 5. RECOMENDAÇÕES ESTRATÉGICAS
Exatamente 5 recomendações objetivas, numeradas, cada uma com o(s) número(s) de processo que a sustenta.
Foco em: teses com maior chance, relatores favoráveis a quais matérias, o que evitar.

Baseie-se exclusivamente nos dados fornecidos. Cite sempre o número do processo ao mencionar um caso."""


def gerar_relatorio(linhas, api_key, model, log):
    """Gera relatório analítico a partir das linhas de resultado processadas."""
    validas = [l for l in linhas
               if l.get("STATUS DA DECISÃO") and l.get("STATUS DA DECISÃO") not in ("", "ERRO", "NÃO LOCALIZADO")]
    if not validas:
        log("   ⚠️ Nenhum processo válido para gerar relatório.")
        return ""

    tabela = []
    for l in validas:
        resumo = (l.get("RESUMO DO PROCESSO") or "")[:300]
        tabela.append(
            f"• Processo: {l.get('NÚMERO DO PROCESSO','')} | "
            f"Data: {l.get('DATA DA DECISÃO','')} | "
            f"Relator: {l.get('RELATOR/JUIZ','')} | "
            f"Turma: {l.get('TURMA/VARA','')} | "
            f"Tipo: {l.get('TIPO','')} | "
            f"Decisão: {l.get('STATUS DA DECISÃO','')} | "
            f"Matéria: {l.get('MATÉRIA','')} | "
            f"Dano Moral: {l.get('DANO MORAL','')} | "
            f"Dano Material: {l.get('DANO MATERIAL','')} | "
            f"Total: {l.get('VALOR DA CONDENAÇÃO','')} | "
            f"Transitado: {l.get('TRANSITADO EM JULGADO? (SIM OU NÃO)','')} | "
            f"Resumo: {resumo}"
        )

    prompt = _PROMPT_RELATORIO_USER.format(n=len(validas), tabela="\n".join(tabela))

    log(f"   📝 Gerando relatório analítico ({len(validas)} processos)...")
    try:
        texto = _chamar_llm(_PROMPT_RELATORIO_SISTEMA, prompt, model, api_key, max_tokens=4000)
        return texto or ""
    except CreditoEsgotadoError:
        # tenta próximo provider
        todas = _carregar_api_keys()
        for _, prov, _ in MODELOS_POR_PROVEDOR:
            chave = todas.get(prov, "").strip()
            if not chave:
                continue
            m = MODELO_PADRAO_POR_PROVIDER[prov]
            if m == model:
                continue
            log(f"   🔄 Sem crédito — relatório com {m}...")
            try:
                return _chamar_llm(_PROMPT_RELATORIO_SISTEMA, prompt, m, chave, max_tokens=4000) or ""
            except CreditoEsgotadoError:
                continue
        log("   ❌ Sem crédito em todos os provedores para gerar relatório.")
        return ""
    except Exception as e:
        if _e_erro_credito(e):
            todas = _carregar_api_keys()
            for _, prov, _ in MODELOS_POR_PROVEDOR:
                chave = todas.get(prov, "").strip()
                if not chave:
                    continue
                m = MODELO_PADRAO_POR_PROVIDER[prov]
                if m == model:
                    continue
                log(f"   🔄 Sem crédito — relatório com {m}...")
                try:
                    return _chamar_llm(_PROMPT_RELATORIO_SISTEMA, prompt, m, chave, max_tokens=4000) or ""
                except Exception:
                    continue
            log("   ❌ Sem crédito em todos os provedores para gerar relatório.")
            return ""
        log(f"   ❌ Erro ao gerar relatório: {e}")
        return ""


def gerar_docx(texto_relatorio, caminho_saida, total_processos=0, nome_advogado=""):
    """Converte o texto do relatório em arquivo DOCX formatado."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import datetime
    except ImportError:
        return False

    # Regex para destacar números de processo (formato CNJ)
    _RE_PROC = re.compile(r'\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}')

    def _add_paragraph_with_proc_highlight(doc, texto, style=None):
        """Adiciona parágrafo com números de processo em negrito."""
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        partes = _RE_PROC.split(texto)
        matches = _RE_PROC.findall(texto)
        for i, parte in enumerate(partes):
            if parte:
                p.add_run(parte)
            if i < len(matches):
                run = p.add_run(matches[i])
                run.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # azul escuro
        return p

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cabeçalho ──────────────────────────────────────────────
    titulo = doc.add_heading("RELATÓRIO ESTRATÉGICO DE PROCESSOS", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("TJAM — Turmas Recursais")
    sub_run.bold = True
    sub_run.font.size = Pt(11)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hoje = datetime.date.today().strftime("%d/%m/%Y")
    info_run = info.add_run(
        f"Gerado em: {hoje}   •   Total analisado: {total_processos} processos"
        + (f"   •   Advogado: {nome_advogado}" if nome_advogado else "")
    )
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Linha separadora
    doc.add_paragraph("─" * 72)

    # ── Conteúdo ───────────────────────────────────────────────
    for linha in texto_relatorio.split("\n"):
        linha_strip = linha.strip()
        if not linha_strip:
            doc.add_paragraph()
            continue

        if linha_strip.startswith("## "):
            texto_sec = linha_strip[3:].strip()
            h = doc.add_heading(texto_sec, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                run.font.size = Pt(12)

        elif linha_strip.startswith("### "):
            texto_sub = linha_strip[4:].strip()
            h = doc.add_heading(texto_sub, level=2)
            for run in h.runs:
                run.font.size = Pt(11)

        elif linha_strip.startswith("▸ ") or linha_strip.startswith("• "):
            _add_paragraph_with_proc_highlight(doc, linha_strip)

        elif linha_strip.startswith(("- ", "* ")):
            p = _add_paragraph_with_proc_highlight(doc, linha_strip[2:], style="List Bullet")

        else:
            _add_paragraph_with_proc_highlight(doc, linha_strip)

    doc.save(caminho_saida)
    return True